#!/usr/bin/env python3
"""
Create Sample DOCX File

A utility script to create a sample DOCX file for testing the knowledge base file parser.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_docx(output_path):
    """
    Create a sample DOCX file with formatted content
    
    Args:
        output_path: Path where the DOCX file will be saved
    """
    document = Document()
    
    # Set document properties
    document.core_properties.title = "Dana AI Knowledge Base Overview"
    document.core_properties.author = "Dana AI Team"
    document.core_properties.category = "Documentation"
    document.core_properties.comments = "Sample document for knowledge base testing"
    document.core_properties.subject = "Knowledge Base"
    
    # Add a title
    title = document.add_heading('Dana AI Knowledge Base System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a subtitle
    subtitle = document.add_paragraph('Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    # Add an introduction
    document.add_heading('Introduction', level=1)
    document.add_paragraph(
        'The Dana AI Platform provides a comprehensive knowledge base system that '
        'allows organizations to store, organize, and retrieve information effectively. '
        'This document outlines the key components and functionality of the system.'
    )
    
    # Add system components section
    document.add_heading('System Components', level=1)
    document.add_paragraph(
        'The knowledge base system consists of the following core components:'
    )
    
    # Add a bulleted list
    components = [
        'File Storage and Management: Handles uploading, versioning, and organization of files.',
        'Content Extraction: Parses various file formats to extract searchable text and metadata.',
        'Search Engine: Provides full-text and semantic search capabilities across all content.',
        'API Interface: Offers RESTful endpoints for programmatic access to knowledge content.',
        'Integration System: Connects with external platforms and services for data exchange.'
    ]
    
    for component in components:
        p = document.add_paragraph(component, style='List Bullet')
    
    # Add file format support section
    document.add_heading('Supported File Formats', level=1)
    
    # Create a table for file formats
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Format'
    header_cells[1].text = 'File Extension'
    header_cells[2].text = 'Features'
    
    # Add data rows
    formats = [
        ('PDF Documents', '.pdf', 'Full text extraction, metadata, page-by-page content'),
        ('Word Documents', '.docx', 'Text extraction, metadata, paragraph-level content'),
        ('Plain Text', '.txt', 'Full text import, line-by-line processing'),
        ('Rich Text Format', '.rtf', 'Formatted text extraction')
    ]
    
    for format_name, extension, features in formats:
        row_cells = table.add_row().cells
        row_cells[0].text = format_name
        row_cells[1].text = extension
        row_cells[2].text = features
    
    # Add integration section
    document.add_heading('Integration Capabilities', level=1)
    
    p = document.add_paragraph(
        'The knowledge base system integrates with various components of the Dana AI Platform, including:'
    )
    
    # Add another bulleted list
    integrations = [
        'Conversation System: Provides relevant information for automated responses.',
        'User Dashboard: Allows users to browse and search knowledge content.',
        'Analytics Engine: Tracks knowledge base usage and effectiveness.',
        'External APIs: Connects with third-party services and data sources.'
    ]
    
    for integration in integrations:
        p = document.add_paragraph(integration, style='List Bullet')
    
    # Add implementation details
    document.add_heading('Implementation Details', level=1)
    document.add_paragraph(
        'The knowledge base system is implemented using the following technologies:'
    )
    
    # Create a table for technology stack
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Technology'
    
    # Add data rows
    technologies = [
        ('Database', 'Supabase (PostgreSQL)'),
        ('File Storage', 'Supabase Storage'),
        ('Search Engine', 'PostgreSQL Full-Text Search'),
        ('API Layer', 'Flask RESTful API'),
        ('File Parsing', 'Custom Python Utilities (PyPDF2, python-docx)')
    ]
    
    for component, technology in technologies:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = technology
    
    # Add security considerations
    document.add_heading('Security Considerations', level=1)
    
    p = document.add_paragraph(
        'The knowledge base system implements several security measures to protect sensitive information:'
    )
    
    security_measures = [
        'User Authentication: All access requires valid authentication credentials.',
        'Row-Level Security: Content is isolated at the database level by user ID.',
        'Access Control: Granular permissions control who can view, edit, or delete content.',
        'Content Validation: Uploaded files are scanned and validated before processing.',
        'Audit Logging: All operations are logged for security monitoring and compliance.'
    ]
    
    for measure in security_measures:
        p = document.add_paragraph(measure, style='List Bullet')
    
    # Add a conclusion
    document.add_heading('Conclusion', level=1)
    document.add_paragraph(
        'The Dana AI Knowledge Base System provides a robust foundation for storing, organizing, '
        'and retrieving information across the platform. It enables organizations to leverage '
        'their existing documentation and knowledge assets to enhance customer interactions, '
        'streamline operations, and improve decision-making processes.'
    )
    
    # Save the document
    document.save(output_path)
    print(f"Sample DOCX file created at: {output_path}")

def main():
    """Main function"""
    output_dir = "sample_docs"
    
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Create the sample DOCX file
    output_path = os.path.join(output_dir, "dana_knowledge_base.docx")
    create_sample_docx(output_path)

if __name__ == "__main__":
    main()